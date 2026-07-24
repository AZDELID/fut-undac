# -*- coding: utf-8 -*-
"""
Generador del documento .docx del FUT UNDAC, en Python puro
(python-docx), sin dependencia de Node.js.

Replica fielmente el formato oficial: logo + título centrado,
"SOLICITO" alineado a la derecha arriba, 14 campos numerados con
líneas punteadas, y el talón/cargo de recepción al final.

Incluye modo compacto adaptativo: si la fundamentación es larga,
reduce automáticamente fuente y espaciados para que el documento
siga cabiendo en una sola página (como el formato físico oficial).
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from . import plantilla as plantilla_mod


def _dir_data() -> Path:
    """Ubica la carpeta data/ con recursos (el logo). En modo ejecutable
    PyInstaller, los recursos empaquetados con --add-data viven en
    sys._MEIPASS (carpeta temporal de extracción del bundle)."""
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return Path(sys._MEIPASS) / "data"
    return Path(__file__).resolve().parent.parent / "data"


DIR_DATA = _dir_data()
LOGO_PATH = DIR_DATA / "undac_logo.jpg"
WATERMARK_PATH = DIR_DATA / "undac_watermark.png"

FONT_NAME = "Arial"


# ---------------------------------------------------------------------
# Marca de agua (escudo UNDAC semitransparente, centrado en la página)
# ---------------------------------------------------------------------

def _insertar_marca_de_agua(doc, section, config: dict = None):
    """Inserta el escudo UNDAC como marca de agua de fondo centrada en la
    página. Usa un drawing anclado con behindDoc=1 en el primer párrafo
    del body, con posición absoluta centrada en la página.

    Se puede desactivar por completo desde la plantilla (config["mostrar_marca_agua"])."""
    if config is not None and not config.get("mostrar_marca_agua", True):
        return
    if not WATERMARK_PATH.exists():
        return

    from docx.shared import Cm
    import copy

    # Agregar imagen temporalmente como párrafo normal para que python-docx
    # registre el archivo en el media del documento y genere el rId
    tmp_p = doc.add_paragraph()
    tmp_run = tmp_p.add_run()
    tmp_run.add_picture(str(WATERMARK_PATH), width=Cm(12), height=Cm(12))

    # Extraer el elemento drawing generado
    NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    NS_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    NS_A  = "http://schemas.openxmlformats.org/drawingml/2006/main"

    drawing_elem = tmp_p._element.find(f"{{{NS_W}}}r/{{{NS_W}}}drawing")
    if drawing_elem is None:
        tmp_p._element.getparent().remove(tmp_p._element)
        return

    drawing_copy = copy.deepcopy(drawing_elem)
    # Eliminar párrafo temporal
    tmp_p._element.getparent().remove(tmp_p._element)

    # Transformar wp:inline -> wp:anchor con behindDoc=1 y posición absoluta
    inline = drawing_copy.find(f"{{{NS_WP}}}inline")
    if inline is None:
        return

    # Dimensiones en EMU (1 cm = 360000 EMU)
    WM_SIZE = int(13.0 * 360000)
    # Centrado horizontal: margen_izq(1.9) + (17.7-13)/2 = 1.9+2.35 = 4.25cm desde borde
    H_POS = str(int(4.05 * 360000))
    # Vertical: zona de campos (entre SUMILLA y FIRMA), ≈7.5cm desde borde superior
    V_POS = str(int(7.2 * 360000))

    anchor = OxmlElement("wp:anchor")
    anchor.set("distT","0"); anchor.set("distB","0")
    anchor.set("distL","0"); anchor.set("distR","0")
    anchor.set("simplePos","0")
    anchor.set("relativeHeight","251658240")
    anchor.set("behindDoc","1")
    anchor.set("locked","1")
    anchor.set("layoutInCell","1")
    anchor.set("allowOverlap","0")

    # Orden correcto según schema CT_Anchor:
    # simplePos, positionH, positionV, extent, effectExtent,
    # wrap*, docPr, cNvGraphicFramePr, a:graphic

    sp = OxmlElement("wp:simplePos"); sp.set("x","0"); sp.set("y","0")
    anchor.append(sp)

    ph = OxmlElement("wp:positionH"); ph.set("relativeFrom","page")
    ph_off = OxmlElement("wp:posOffset"); ph_off.text = H_POS
    ph.append(ph_off); anchor.append(ph)

    pv = OxmlElement("wp:positionV"); pv.set("relativeFrom","page")
    pv_off = OxmlElement("wp:posOffset"); pv_off.text = V_POS
    pv.append(pv_off); anchor.append(pv)

    # Copiar extent desde el inline (contiene las dimensiones correctas en EMU)
    extent = inline.find(f"{{{NS_WP}}}extent")
    if extent is not None:
        anchor.append(copy.deepcopy(extent))

    ee = OxmlElement("wp:effectExtent")
    ee.set("l","0"); ee.set("t","0"); ee.set("r","0"); ee.set("b","0")
    anchor.append(ee)

    # wrapNone ANTES de docPr (orden del schema)
    anchor.append(OxmlElement("wp:wrapNone"))

    # docPr
    dp = inline.find(f"{{{NS_WP}}}docPr")
    if dp is not None:
        anchor.append(copy.deepcopy(dp))

    # cNvGraphicFramePr
    cg = inline.find(f"{{{NS_WP}}}cNvGraphicFramePr")
    if cg is not None:
        anchor.append(copy.deepcopy(cg))

    # a:graphic (el bloque de imagen real)
    graphic = inline.find(f"{{{NS_A}}}graphic")
    if graphic is not None:
        anchor.append(copy.deepcopy(graphic))

    # Reemplazar inline por anchor
    drawing_copy.remove(inline)
    drawing_copy.append(anchor)

    # Insertar como PRIMER párrafo del body para que esté detrás de todo
    wm_p = OxmlElement("w:p")
    wm_ppr = OxmlElement("w:pPr")
    wm_p.append(wm_ppr)
    wm_r = OxmlElement("w:r")
    wm_p.append(wm_r)
    wm_r.append(drawing_copy)

    # Insertar al inicio del body (antes del primer elemento)
    body = doc.element.body
    body.insert(0, wm_p)


# ---------------------------------------------------------------------
# Helpers de bajo nivel (bordes punteados, anchos de tabla, etc.)
# ---------------------------------------------------------------------

def _set_cell_borders_none(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def _add_dotted_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "dotted")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_borders.append(bottom)
    # El schema OOXML exige un orden estricto de hijos de pPr; pBdr debe
    # insertarse antes que spacing/ind/tabs/jc, no simplemente al final.
    _insert_pPr_child_ordered(p_pr, p_borders, "pBdr")


# Orden canónico (subconjunto relevante) de CT_PPrBase según el schema OOXML
_PPR_ORDER = [
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr",
    "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs",
    "suppressAutoHyphens", "kinsoku", "wordWrap", "overflowPunct",
    "topLinePunct", "autoSpaceDE", "autoSpaceDN", "bidi", "adjustRightInd",
    "snapToGrid", "spacing", "ind", "contextualSpacing", "mirrorIndents",
    "suppressOverlap", "jc", "textDirection", "textAlignment",
    "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr",
]


def _insert_pPr_child_ordered(p_pr, element, tag_name):
    target_idx = _PPR_ORDER.index(tag_name)
    insert_pos = len(p_pr)
    for i, child in enumerate(p_pr):
        local_tag = child.tag.split("}")[-1]
        if local_tag in _PPR_ORDER and _PPR_ORDER.index(local_tag) > target_idx:
            insert_pos = i
            break
    p_pr.insert(insert_pos, element)


def _set_paragraph_spacing(paragraph, before=0, after=6):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)


def _run(paragraph, text, bold=False, size=10.5, color=None):
    r = paragraph.add_run(text if text else " ")
    r.bold = bold
    r.font.name = FONT_NAME
    r.font.size = Pt(size)
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r


def _set_page_margins(section, top_cm, bottom_cm, side_cm=1.9):
    section.top_margin = Cm(top_cm)
    section.bottom_margin = Cm(bottom_cm)
    section.left_margin = Cm(side_cm)
    section.right_margin = Cm(side_cm)


# ---------------------------------------------------------------------
# Wrap de texto para simular líneas punteadas tipo formulario impreso
# ---------------------------------------------------------------------

def _wrap_texto(texto, ancho):
    palabras = texto.split(" ")
    lineas = []
    actual = ""
    for palabra in palabras:
        candidato = (actual + " " + palabra).strip()
        if len(candidato) > ancho and actual:
            lineas.append(actual)
            actual = palabra
        else:
            actual = candidato
    if actual:
        lineas.append(actual)
    return lineas or [""]


def _wrap_fundamentacion(texto, ancho, min_lineas):
    lineas_originales = [l for l in (texto or "").split("\n") if l.strip() != ""]
    wrapped = []
    for linea in lineas_originales:
        wrapped.extend(_wrap_texto(linea, ancho))
    while len(wrapped) < min_lineas:
        wrapped.append("")
    return wrapped


# ---------------------------------------------------------------------
# Generador principal
# ---------------------------------------------------------------------

def _insertar_pie_de_pagina(doc, section, config: dict):
    """Escribe el pie de página real del documento (dirección + año/período
    + texto libre), usando la configuración de plantilla del usuario."""
    if not config.get("mostrar_pie", True):
        return

    lema = (config.get("lema_anio") or "").strip()
    partes = [p.strip() for p in (
        lema if config.get("mostrar_lema_anio", True) and lema else "",
        config.get("direccion", ""),
        f"Año Académico {config['anio_periodo']}" if config.get("anio_periodo") else "",
        config.get("texto_pie", ""),
    ) if p and p.strip()]
    if not partes:
        return

    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    color_pie = (config.get("color_pie") or "6b7280").strip().lstrip("#") or "6b7280"
    _run(p, "  •  ".join(partes), size=7.5, color=color_pie)


def generar_docx(fut_dict: dict, ruta_salida: str):
    """
    fut_dict: diccionario (FUTData.to_dict()) con los datos del FUT.
    ruta_salida: ruta donde se guardará el .docx
    """
    config_plantilla = plantilla_mod.obtener_plantilla()

    largo_fundamentacion = len(fut_dict.get("fundamentacion") or "")
    compacto = largo_fundamentacion > 380

    tam_base = 9.5 if compacto else 10.5
    tam_talon = 8 if compacto else 8.5
    espacio_after_dotted = 5 if compacto else 7
    espacio_campo_before = 3 if compacto else 4
    min_lineas_fund = 3 if compacto else 5
    ancho_wrap = 100 if compacto else 92
    margen_top = 0.9 if compacto else 1.1
    margen_bottom = 0.9 if compacto else 1.1

    doc = Document()
    section = doc.sections[0]
    _set_page_margins(section, margen_top, margen_bottom)

    # Estilo normal por defecto
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(tam_base)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    # ---------------- Encabezado: logo izq. + título + logo escuela ------
    mostrar_logo_izq = config_plantilla.get("mostrar_logo_undac", True)
    mostrar_logo_der = bool(
        config_plantilla.get("mostrar_logo_escuela", False)
        and config_plantilla.get("logo_escuela_path", "")
    )
    ruta_logo_escuela = Path(config_plantilla.get("logo_escuela_path", "") or "")

    tabla_header = doc.add_table(rows=1, cols=3)
    tabla_header.autofit = False
    tabla_header.columns[0].width = Cm(2.2)
    tabla_header.columns[1].width = Cm(12.3)
    tabla_header.columns[2].width = Cm(2.2)

    # -- Logo institucional (izquierda) --
    celda_logo = tabla_header.rows[0].cells[0]
    celda_logo.width = Cm(2.2)
    _set_cell_borders_none(celda_logo)
    celda_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = celda_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if mostrar_logo_izq and LOGO_PATH.exists():
        run_logo = p_logo.add_run()
        run_logo.add_picture(str(LOGO_PATH), width=Cm(1.7), height=Cm(2.0))

    # -- Título central --
    celda_titulo = tabla_header.rows[0].cells[1]
    _set_cell_borders_none(celda_titulo)
    celda_titulo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    encabezado_extra = (config_plantilla.get("encabezado_personalizado") or "").strip()
    p_titulo1 = celda_titulo.paragraphs[0]
    if encabezado_extra:
        p_titulo1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p_titulo1, before=0, after=1)
        _run(p_titulo1, encabezado_extra, bold=False, size=9)
        p_titulo1 = celda_titulo.add_paragraph()
    p_titulo1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p_titulo1, before=0, after=0)
    _run(p_titulo1, config_plantilla["nombre_institucion"], bold=True, size=13.5)
    p_titulo2 = celda_titulo.add_paragraph()
    p_titulo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p_titulo2, before=0, after=0)
    _run(p_titulo2, config_plantilla["subtitulo"], bold=True, size=12)

    # -- Logo de la escuela/facultad (derecha, opcional) --
    celda_logo_der = tabla_header.rows[0].cells[2]
    celda_logo_der.width = Cm(2.2)
    _set_cell_borders_none(celda_logo_der)
    celda_logo_der.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo_der = celda_logo_der.paragraphs[0]
    p_logo_der.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if mostrar_logo_der and ruta_logo_escuela.exists():
        try:
            run_logo_der = p_logo_der.add_run()
            run_logo_der.add_picture(str(ruta_logo_escuela), width=Cm(1.7), height=Cm(1.7))
        except Exception:
            pass

    # Línea gruesa separadora
    p_linea = doc.add_paragraph()
    _set_paragraph_spacing(p_linea, before=4, after=10)
    p_pr = p_linea._p.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_borders.append(bottom)
    _insert_pPr_child_ordered(p_pr, p_borders, "pBdr")

    # ---------------- Nº de expediente/folio (opcional) ----------------
    if config_plantilla.get("mostrar_numero_expediente", False):
        prefijo = (config_plantilla.get("prefijo_expediente") or "N° EXP.").strip()
        p_exp = doc.add_paragraph()
        p_exp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_paragraph_spacing(p_exp, before=0, after=6)
        _run(p_exp, f"{prefijo}: ______________", size=9, color="4b5563")

    # ---------------- Datos del estudiante en el encabezado (opcional) --
    if config_plantilla.get("mostrar_datos_estudiante", False):
        nombre_est = (fut_dict.get("nombres_apellidos") or "").strip()
        codigo_est = (fut_dict.get("codigo") or "").strip()
        if nombre_est or codigo_est:
            partes_est = [p for p in (
                f"Estudiante: {nombre_est}" if nombre_est else "",
                f"Código: {codigo_est}" if codigo_est else "",
            ) if p]
            p_est = doc.add_paragraph()
            p_est.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_paragraph_spacing(p_est, before=0, after=6)
            _run(p_est, "  •  ".join(partes_est), size=9, color="4b5563")

    # ---------------- Bloque SOLICITO ----------------
    solicito_lineas = _wrap_texto(fut_dict.get("solicito", ""), 38)
    p_sol1 = doc.add_paragraph()
    p_sol1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_dotted_bottom_border(p_sol1)
    _set_paragraph_spacing(p_sol1, before=0, after=4)
    _run(p_sol1, "SOLICITO: ", bold=True, size=tam_base)
    _run(p_sol1, solicito_lineas[0], size=tam_base)
    for linea_extra in solicito_lineas[1:3]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_dotted_bottom_border(p)
        _set_paragraph_spacing(p, before=0, after=4)
        _run(p, linea_extra, size=tam_base)
    # mínimo 3 líneas en el bloque solicito, como el original
    for _ in range(3 - len(solicito_lineas[:3])):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_dotted_bottom_border(p)
        _set_paragraph_spacing(p, before=0, after=4)
        _run(p, "", size=tam_base)

    # ---------------- Helpers de campo numerado ----------------

    def campo_numerado(numero, etiqueta):
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=espacio_campo_before * 2, after=2)
        _run(p, f"{numero}. ", bold=True, size=tam_base)
        _run(p, etiqueta, bold=True, size=tam_base)
        return p

    def campo_numerado_doble(n1, e1, n2, e2):
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=espacio_campo_before * 2, after=2)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(9.5))
        _run(p, f"{n1}. ", bold=True, size=tam_base)
        _run(p, e1, bold=True, size=tam_base)
        _run(p, "\t")
        _run(p, f"{n2}. ", bold=True, size=tam_base)
        _run(p, e2, bold=True, size=tam_base)
        return p

    def linea_dotted(texto):
        p = doc.add_paragraph()
        _add_dotted_bottom_border(p)
        _set_paragraph_spacing(p, before=0, after=espacio_after_dotted)
        _run(p, texto or "", size=tam_base)
        return p

    def linea_dotted_doble(v1, v2):
        p = doc.add_paragraph()
        _add_dotted_bottom_border(p)
        _set_paragraph_spacing(p, before=0, after=espacio_after_dotted)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(9.5))
        _run(p, v1 or " ", size=tam_base)
        _run(p, "\t")
        _run(p, v2 or " ", size=tam_base)
        return p

    # ---------------- Campos 1-10 ----------------
    campo_numerado(1, "SUMILLA")
    linea_dotted(fut_dict.get("sumilla", ""))

    campo_numerado(2, "DESTINATARIO")
    linea_dotted(fut_dict.get("destinatario", ""))

    campo_numerado(3, "DATOS DEL USUARIO (APELLIDOS Y NOMBRES)")
    linea_dotted(fut_dict.get("nombres_apellidos", ""))

    campo_numerado(4, "CARGO ACTUAL Y/O CENTRO DE TRABAJO")
    linea_dotted(fut_dict.get("cargo_centro_trabajo", ""))

    campo_numerado_doble(5, "D.N.I.", 6, "CÓDIGO DE MATRÍCULA")
    linea_dotted_doble(fut_dict.get("dni", ""), fut_dict.get("codigo", ""))

    campo_numerado_doble(7, "N° CELULAR/TELF.", 8, "CORREO ELECTRÓNICO")
    linea_dotted_doble(fut_dict.get("celular", ""), fut_dict.get("correo", ""))

    p9 = doc.add_paragraph()
    _set_paragraph_spacing(p9, before=espacio_campo_before * 2, after=2)
    _run(p9, "9. ", bold=True, size=tam_base)
    _run(p9, "FACULTAD  /  ESCUELA PROFESIONAL  /  ESPECIALIDAD", bold=True, size=tam_base - 1)
    valores9 = [v for v in [fut_dict.get("facultad", ""), fut_dict.get("escuela", ""), fut_dict.get("especialidad", "")] if v and v.strip()]
    linea_dotted("   /   ".join(valores9))

    campo_numerado(10, "DOMICILIO DEL USUARIO (Calle, Distrito, Provincia Y Región)")
    linea_dotted(fut_dict.get("domicilio", ""))

    # ---------------- Fundamentación ----------------
    campo_numerado(11, "FUNDAMENTACIÓN DEL PEDIDO")
    for linea in _wrap_fundamentacion(fut_dict.get("fundamentacion", ""), ancho_wrap, min_lineas_fund):
        linea_dotted(linea)

    # ---------------- Anexo / Fecha / Firma ----------------
    p12 = doc.add_paragraph()
    _set_paragraph_spacing(p12, before=espacio_campo_before * 2, after=2)
    tab_stops = p12.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(9.5))
    _run(p12, "12. ANEXO.", bold=True, size=tam_base)
    _run(p12, "\t")
    if config_plantilla.get("mostrar_fecha_lugar", False):
        lugar_txt = (fut_dict.get("lugar") or config_plantilla.get("lugar_predeterminado", "")).strip()
        etiqueta_fecha = f"13. FECHA Y LUGAR: {lugar_txt}, {fut_dict.get('fecha', '')}" if lugar_txt else f"13. FECHA: {fut_dict.get('fecha', '')}"
    else:
        etiqueta_fecha = f"13. FECHA: {fut_dict.get('fecha', '')}"
    _run(p12, etiqueta_fecha, bold=True, size=tam_base)
    linea_dotted(fut_dict.get("anexo", ""))

    p14 = doc.add_paragraph()
    _set_paragraph_spacing(p14, before=4, after=2)
    _run(p14, "14. FIRMA:", bold=True, size=tam_base)
    linea_dotted("")

    # ---------------- Línea separadora (talón) ----------------
    p_sep = doc.add_paragraph()
    _set_paragraph_spacing(p_sep, before=6, after=6)
    _run(p_sep, "=" * 98, size=8)

    # ---------------- Talón / cargo de recepción ----------------
    p_t1 = doc.add_paragraph()
    _set_paragraph_spacing(p_t1, before=0, after=5)
    _run(p_t1, "FUNDAMENTACIÓN DEL PEDIDO: ", bold=True, size=tam_talon)
    resumen = (fut_dict.get("sumilla") or "")[: (70 if compacto else 90)]
    _run(p_t1, resumen, size=tam_talon)

    p_t2 = doc.add_paragraph()
    _set_paragraph_spacing(p_t2, before=0, after=5)
    tab_stops = p_t2.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(12.5))
    _run(p_t2, "APELLIDOS Y NOMBRES: ", bold=True, size=tam_talon)
    _run(p_t2, fut_dict.get("nombres_apellidos", ""), size=tam_talon)
    _run(p_t2, "\t")
    _run(p_t2, "FOLIO: ", bold=True, size=tam_talon)
    _run(p_t2, "________________", size=tam_talon)

    p_t3 = doc.add_paragraph()
    _set_paragraph_spacing(p_t3, before=0, after=5)
    tab_stops = p_t3.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(12.5))
    _run(p_t3, "N° DE REGISTRO: ", bold=True, size=tam_talon)
    _run(p_t3, "________________", size=tam_talon)
    _run(p_t3, "\t")
    _run(p_t3, "FECHA: ", bold=True, size=tam_talon)
    _run(p_t3, "________________", size=tam_talon)

    p_t4 = doc.add_paragraph()
    _add_dotted_bottom_border(p_t4)
    _set_paragraph_spacing(p_t4, before=0, after=3)
    _run(p_t4, " ", size=tam_talon)

    p_t5 = doc.add_paragraph()
    _set_paragraph_spacing(p_t5, before=0, after=0)
    _run(p_t5, "FACULTAD / ESCUELA PROFESIONAL / ESPECIALIDAD", bold=True, size=tam_talon - 1)

    # ---------------- Marca de agua (escudo UNDAC de fondo) ----------------
    _insertar_marca_de_agua(doc, section, config_plantilla)

    # ---------------- Pie de página (dirección, año, texto libre) --------
    _insertar_pie_de_pagina(doc, section, config_plantilla)

    doc.save(ruta_salida)
    _corregir_zoom_settings(ruta_salida)
    return ruta_salida


def _corregir_zoom_settings(ruta_docx):
    """python-docx a veces deja <w:zoom w:val="bestFit"/> sin el atributo
    w:percent requerido por el schema OOXML cuando se insertan imágenes.
    Esta función agrega el atributo faltante post-guardado."""
    import zipfile
    import shutil
    import tempfile
    import re
    import os

    with zipfile.ZipFile(ruta_docx, "r") as zin:
        if "word/settings.xml" not in zin.namelist():
            return
        settings_xml = zin.read("word/settings.xml").decode("utf-8")

    if 'w:zoom' not in settings_xml:
        return

    nuevo_xml, n = re.subn(
        r'(<w:zoom\b(?:(?!w:percent)[^/])*?)/>',
        r'\1 w:percent="100"/>',
        settings_xml,
    )
    if n == 0:
        return

    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(tmp_fd)
    with zipfile.ZipFile(ruta_docx, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == "word/settings.xml":
                zout.writestr(item, nuevo_xml)
            else:
                zout.writestr(item, zin.read(item.filename))
    shutil.move(tmp_path, ruta_docx)
